"""Export service – generates PDF and DOCX files from note data.

Reuses the existing Note model and service layer. The generated files are
returned as in-memory byte streams ready for HTTP streaming.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.note import Note
from app.services.font_manager import register_fonts, sanitize_text
from fpdf import FPDF

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────
#  PDF generation
# ──────────────────────────────────────────────────────────────────────


class NotesPDFBase(FPDF):
    """Minimal FPDF subclass that reads ``_font_family`` from the instance.

    We store the resolved font family on the *instance* so that concurrent
    PDF generation (via thread-pool executor) does not share stale global
    state — each thread calls :func:`register_fonts` independently.
    """

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._font_family: str = "Helvetica"  # set during generate_pdf()


class NotesPDF(NotesPDFBase):
    """Clean, production-grade PDF layout for exported notes."""

    def header(self) -> None:
        if self.page_no() <= 1:
            return
        self.set_font(self._font_family, "I", 8)
        self.set_text_color(140, 140, 140)
        self.cell(0, 8, "NoteFlow AI \u2014 Generated Notes", align="C")
        self.ln(10)

    def footer(self) -> None:
        self.set_y(-15)
        self.set_font(self._font_family, "I", 8)
        self.set_text_color(160, 160, 160)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def _section_heading(pdf: NotesPDF, title: str) -> None:
    """Render a section heading with an underline."""
    pdf.set_font(pdf._font_family, "B", 13)
    pdf.set_text_color(40, 40, 40)
    pdf.cell(0, 8, title)
    pdf.ln(3)
    # Thin underline
    pdf.set_draw_color(200, 200, 200)
    pdf.set_line_width(0.4)
    x = pdf.get_x()
    y = pdf.get_y()
    pdf.line(x, y, x + 190, y)
    pdf.ln(6)


def _body_text(pdf: NotesPDF, text: str) -> None:
    """Render a paragraph of body text with word wrapping."""
    pdf.set_font(pdf._font_family, "", 10)
    pdf.set_text_color(60, 60, 60)
    pdf.multi_cell(0, 5.5, sanitize_text(text))
    pdf.ln(2)


def _bullet_char(pdf: NotesPDF) -> str:
    """Return a bullet character safe for the active font.

    DejaVuSans (and other Unicode TTF fonts) support U+2022 BULLET.
    Helvetica (Latin-1 only) does not — we fall back to a simple
    ASCII hyphen-minus.  This is a **final fallback** only when no
    Unicode font is available.
    """
    if pdf._font_family and pdf._font_family != "Helvetica":
        return "\u2022"  # •
    return "-"


def _bullet_list(pdf: NotesPDF, items: list[str]) -> None:
    """Render a list of bullet points."""
    pdf.set_font(pdf._font_family, "", 10)
    pdf.set_text_color(60, 60, 60)
    bullet = _bullet_char(pdf)
    for item in items:
        pdf.set_x(15)
        pdf.cell(5, 5.5, bullet)
        pdf.multi_cell(0, 5.5, sanitize_text(item))
        pdf.ln(1)
    pdf.ln(2)


def _tag_row(pdf: NotesPDF, tags: list[str]) -> None:
    """Render inline tags / keywords as a single line."""
    pdf.set_font(pdf._font_family, "", 10)
    pdf.set_text_color(60, 60, 60)
    line = "  ".join(tags)
    pdf.multi_cell(0, 5.5, sanitize_text(line))
    pdf.ln(2)


def generate_pdf(note: Note) -> io.BytesIO:
    """Generate a professional PDF from a :class:`Note` model instance.

    Returns an in-memory bytes buffer ready to be streamed as an HTTP
    response.
    """
    pdf = NotesPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=22)

    # ── Register Unicode font (DejaVuSans bundled or system Noto) ───
    family = register_fonts(pdf)
    pdf._font_family = family

    # ── Page 1: Title block ──────────────────────────────────────────
    pdf.add_page()
    pdf.ln(10)
    pdf.set_font(family, "B", 22)
    pdf.set_text_color(25, 25, 25)
    pdf.multi_cell(0, 10, sanitize_text(note.title or "Untitled Notes"), align="C")
    pdf.ln(4)

    pdf.set_font(family, "", 9)
    pdf.set_text_color(140, 140, 140)
    pdf.cell(0, 5, f"Generated on {datetime.now().strftime('%B %d, %Y')}", align="C")
    pdf.ln(4)
    if note.model_used:
        pdf.cell(0, 5, f"Model: {note.model_used}", align="C")
        pdf.ln(4)
    if note.processing_time and note.processing_time > 0:
        pdf.cell(0, 5, f"Processing time: {note.processing_time:.1f}s", align="C")
    pdf.ln(12)

    # ── 1. Executive Summary ─────────────────────────────────────────
    summary = note.executive_summary or ""
    if summary.strip():
        _section_heading(pdf, "Executive Summary")
        _body_text(pdf, summary)

    # ── 2. Key Concepts ──────────────────────────────────────────────
    concepts: list[str] = note.key_concepts or []
    if concepts:
        _section_heading(pdf, "Key Concepts")
        _bullet_list(pdf, concepts)

    # ── 3. Detailed Notes ────────────────────────────────────────────
    detailed = note.detailed_notes or ""
    if detailed.strip():
        _section_heading(pdf, "Detailed Notes")
        _body_text(pdf, detailed)

    # ── 4. Key Takeaways (Bullet Points) ─────────────────────────────
    bullets: list[str] = note.bullet_points or []
    if bullets:
        _section_heading(pdf, "Key Takeaways")
        _bullet_list(pdf, bullets)

    # ── 5. Keywords ──────────────────────────────────────────────────
    keywords: list[str] = note.keywords or []
    if keywords:
        _section_heading(pdf, "Keywords")
        _tag_row(pdf, keywords)

    # ── 6. Action Items ──────────────────────────────────────────────
    actions: list[str] = note.action_items or []
    if actions:
        _section_heading(pdf, "Action Items")
        _bullet_list(pdf, actions)

    # ── 7. Conclusion ────────────────────────────────────────────────
    conclusion = note.conclusion or ""
    if conclusion.strip():
        _section_heading(pdf, "Conclusion")
        _body_text(pdf, conclusion)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


# ──────────────────────────────────────────────────────────────────────
#  DOCX generation
# ──────────────────────────────────────────────────────────────────────


def _docx_set_style(doc: Document) -> None:
    """Tweak the built-in Normal and Heading styles."""
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        try:
            h = doc.styles[f"Heading {level}"]
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0x28, 0x28, 0x28)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
        except KeyError:
            pass


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def generate_docx(note: Note) -> io.BytesIO:
    """Generate a Word document from a :class:`Note` model instance."""
    doc = Document()
    _docx_set_style(doc)

    # ── Title page info ──────────────────────────────────────────────
    _add_heading(doc, note.title or "Untitled Notes", level=1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if note.model_used:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(f"Model: {note.model_used}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if note.processing_time and note.processing_time > 0:
        p3 = doc.add_paragraph()
        run3 = p3.add_run(f"Processing time: {note.processing_time:.1f}s")
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── 1. Executive Summary ─────────────────────────────────────────
    summary = (note.executive_summary or "").strip()
    if summary:
        _add_heading(doc, "Executive Summary", level=2)
        _add_paragraph(doc, summary)

    # ── 2. Key Concepts ──────────────────────────────────────────────
    concepts: list[str] = note.key_concepts or []
    if concepts:
        _add_heading(doc, "Key Concepts", level=2)
        _add_bullet_list(doc, concepts)

    # ── 3. Detailed Notes ────────────────────────────────────────────
    detailed = (note.detailed_notes or "").strip()
    if detailed:
        _add_heading(doc, "Detailed Notes", level=2)
        _add_paragraph(doc, detailed)

    # ── 4. Key Takeaways (Bullet Points) ─────────────────────────────
    bullets: list[str] = note.bullet_points or []
    if bullets:
        _add_heading(doc, "Key Takeaways", level=2)
        _add_bullet_list(doc, bullets)

    # ── 5. Keywords ──────────────────────────────────────────────────
    keywords: list[str] = note.keywords or []
    if keywords:
        _add_heading(doc, "Keywords", level=2)
        _add_paragraph(doc, ", ".join(keywords))

    # ── 6. Action Items ──────────────────────────────────────────────
    actions: list[str] = note.action_items or []
    if actions:
        _add_heading(doc, "Action Items", level=2)
        _add_bullet_list(doc, actions)

    # ── 7. Conclusion ────────────────────────────────────────────────
    conclusion = (note.conclusion or "").strip()
    if conclusion:
        _add_heading(doc, "Conclusion", level=2)
        _add_paragraph(doc, conclusion)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
